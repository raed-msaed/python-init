from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)



document = Document()
# profile picture

document.add_picture(
    'images.jpg',
     width=Inches(2.0)
)

name = input('What is your name? ')
speak('Hello ' + name + ' how are you today')
speak('What is your phone number ? ')
phone_number =input('What is your phone number ? ')
email = input('What is your email? ')

document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

document.add_heading('About me')
document.add_paragraph(input('Tell about yourself? '))
# Work experience
document.add_heading('Work Experience? ')
p = document.add_paragraph()
company = input('inserer your company? ')
p.add_run(company)
document.save('cv.docx')